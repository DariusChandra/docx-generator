import tkinter as tk
from tkinter import filedialog
from docx import Document
import pandas as pd
import os

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("DOCX Generator")
        self.geometry("400x200")

        # Create buttons to select excel and docx files
        self.excel_button = tk.Button(self, text="Select Excel File", command=self.open_excel)
        self.excel_button.pack()
        self.docx_button = tk.Button(self, text="Select Template DOCX", command=self.open_docx)
        self.docx_button.pack()

        # Create a button to start the DOCX generation process
        self.generate_button = tk.Button(self, text="Generate DOCXs", command=self.generate_docxs)
        self.generate_button.pack()

        # Create a label to display the status of the DOCX generation process
        self.status_label = tk.Label(self, text="")
        self.status_label.pack()

    def open_excel(self):
        self.excel_file = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        self.status_label.config(text=f"Excel file selected: {self.excel_file}")

    def open_docx(self):
        self.docx_file = filedialog.askopenfilename(filetypes=[("DOCX files", "*.docx")])
        self.status_label.config(text=f"DOCX file selected: {self.docx_file}")

    def generate_docxs(self):
        # Read the data from the excel file
        data = pd.read_excel(self.excel_file)

        # Open the template docx file
        template_doc = Document(self.docx_file)

        for i in range(len(data)):
            # Create a new docx file for each row in the excel file
            doc = Document(self.docx_file)

            # Replace the placeholders with the values from the excel file
            for p in doc.paragraphs:
                for key in data.keys():
                    if f"$({key})" in p.text:
                        p.text = p.text.replace(f"$({key})", str(data[key][i]))

            # check if the file already exist
            docx_file_name = data["doc_name"][i] + '.docx'
            if os.path.exists(docx_file_name):
                os.remove(docx_file_name)
            # Save the new docx file with the name from the excel file
            doc.save(data["doc_name"][i] + '.docx')

        self.status_label.config(text="DOCXs generated successfully!")

if __name__ == "__main__":
    app = App()
    app.mainloop()

